from docx import Document
import os

doc_path = r"d:\Projects\Dify\word_service\test_output_en.docx"

if os.path.exists(doc_path):
    doc = Document(doc_path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    
    print("\n--- Paragraph Texts ---")
    for i, p in enumerate(doc.paragraphs[:20]):
        print(f"{i}: {p.text}")
    
    print("\n--- Tables Content ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
        if i == len(doc.tables) - 1: # Usually the one we inserted
            for row in table.rows[:3]:
                print(f"  Row: {[cell.text for cell in row.cells]}")
else:
    print(f"File not found: {doc_path}")
