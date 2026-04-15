import base64
import io
import unittest
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient

import main


def make_template_base64(rows: int = 6, cols: int = 5) -> str:
    doc = Document()
    doc.add_table(rows=rows, cols=cols)
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode()


def doc_from_base64(document_base64: str) -> Document:
    return Document(io.BytesIO(base64.b64decode(document_base64)))


class GenerateEndpointBehaviorTests(unittest.TestCase):
    def setUp(self):
        self.client = TestClient(main.app)

    def test_generate_updates_cn_en_water_level(self):
        template = make_template_base64()
        with patch("main.fetch_feishu_water_level", return_value="123.45"):
            response = self.client.post(
                "/generate-from-template",
                json={
                    "chinese_data": '[{"seq":1,"location":"右岸道路","content":"测试内容","quantity":"100m"}]',
                    "english_data": '{"translated_data":[{"seq":1,"location_en":"Right Bank Roads","content_en":"Test content","quantity_en":"100m","remarks_en":""}]}',
                    "cn_template_base64": template,
                    "en_template_base64": template,
                    "update_date_weather": True,
                    "feishu_app_token": "app_token",
                    "feishu_table_id": "table_id",
                },
            )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertEqual(body["weather_info"]["water_level"], "123.45")
        self.assertEqual(body["weather_info"]["water_level_status"], "ok")
        self.assertNotIn("warnings", body)

        cn_doc = doc_from_base64(body["cn_document_base64"])
        en_doc = doc_from_base64(body["en_document_base64"])
        self.assertEqual(cn_doc.tables[0].rows[0].cells[3].text, "123.45")
        self.assertEqual(en_doc.tables[0].rows[0].cells[3].text, "123.45")

    def test_generate_returns_warnings_when_feishu_fails(self):
        template = make_template_base64()
        with patch("main.fetch_feishu_water_level", side_effect=RuntimeError("boom")):
            response = self.client.post(
                "/generate-from-template",
                json={
                    "chinese_data": '[{"seq":1,"location":"右岸道路","content":"测试内容","quantity":"100m"}]',
                    "english_data": '{"translated_data":[{"seq":1,"location_en":"Right Bank Roads","content_en":"Test content","quantity_en":"100m","remarks_en":""}]}',
                    "cn_template_base64": template,
                    "en_template_base64": template,
                    "update_date_weather": True,
                    "feishu_app_token": "app_token",
                    "feishu_table_id": "table_id",
                },
            )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        self.assertIn("warnings", body)
        self.assertEqual(len(body["warnings"]), 2)
        self.assertEqual(body["weather_info"]["water_level"], "")
        self.assertEqual(body["weather_info"]["water_level_status"], "failed")
        self.assertIn("cn_document_base64", body)
        self.assertIn("en_document_base64", body)

    def test_cn_en_use_same_grouped_table_logic(self):
        doc = Document()
        table = doc.add_table(rows=6, cols=6)
        table.cell(0, 0).text = "序号"
        table.cell(1, 0).text = "1"
        table.cell(3, 0).text = "2"
        table.cell(1, 1).merge(table.cell(2, 1))
        table.cell(3, 1).merge(table.cell(4, 1))
        buf = io.BytesIO()
        doc.save(buf)
        template = base64.b64encode(buf.getvalue()).decode()

        response = self.client.post(
            "/generate-from-template",
            json={
                "chinese_data": [
                    {"seq": 1, "location": "右岸道路", "content": "内容A", "quantity": "10m"},
                    {"seq": 1, "location": "右岸道路", "content": "内容B", "quantity": "20m"},
                    {"seq": 2, "location": "炸药库", "content": "内容C", "quantity": "30m"},
                ],
                "english_data": {
                    "translated_data": [
                        {"seq": 1, "location_en": "Right Bank Roads", "content_en": "Content A", "quantity_en": "10m", "remarks_en": ""},
                        {"seq": 1, "location_en": "Right Bank Roads", "content_en": "Content B", "quantity_en": "20m", "remarks_en": ""},
                        {"seq": 2, "location_en": "Explosives Magazine", "content_en": "Content C", "quantity_en": "30m", "remarks_en": ""},
                    ]
                },
                "cn_template_base64": template,
                "en_template_base64": template,
                "cn_table_index": 0,
                "en_table_index": 0,
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertTrue(body["success"])
        cn_doc = doc_from_base64(body["cn_document_base64"])
        en_doc = doc_from_base64(body["en_document_base64"])

        self.assertEqual(cn_doc.tables[0].rows[1].cells[1].text, "右岸道路")
        self.assertEqual(cn_doc.tables[0].rows[1].cells[3].text, "内容A")
        self.assertEqual(cn_doc.tables[0].rows[2].cells[3].text, "内容B")
        self.assertEqual(en_doc.tables[0].rows[1].cells[1].text, "Right Bank Roads")
        self.assertEqual(en_doc.tables[0].rows[1].cells[3].text, "Content A")
        self.assertEqual(en_doc.tables[0].rows[2].cells[3].text, "Content B")


if __name__ == "__main__":
    unittest.main()
