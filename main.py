"""
Word Document Generation Service - Precision Formatting Version (v8.0)
修复重点：
1. 消除大标题(1、)的错误缩进。
2. 实现"人员投入"等关键词的【局部加粗】（而非整行）。
3. 严格控制子标题((1))和统计项的首行缩进。
"""

import base64
import io
import json
import re
from datetime import datetime
from typing import List, Optional, Dict, Any

import requests
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

app = FastAPI(title="Word Service Precision", version="8.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------- 核心排版逻辑 ----------------

def format_run_font(run, size=10.5, bold=False):
    """统一设置字体格式"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold


def encode_document(doc: Document) -> str:
    out = io.BytesIO()
    doc.save(out)
    return base64.b64encode(out.getvalue()).decode()


def decode_document(document_base64: str) -> Document:
    file_bytes = base64.b64decode(document_base64)
    return Document(io.BytesIO(file_bytes))


def document_response(doc: Document, **extra: Any) -> dict[str, Any]:
    result: dict[str, Any] = {"success": True, "document_base64": encode_document(doc)}
    result.update(extra)
    return result


def update_document(document_base64: str, mutator):
    doc = decode_document(document_base64)
    extra = mutator(doc) or {}
    return document_response(doc, **extra)


def parse_payload(payload: Any) -> list[dict]:
    """兼容 JSON 字符串、dict 包装和 list 三种常见输入。"""
    if isinstance(payload, str):
        payload = payload.strip()
        if not payload:
            return []
        payload = json.loads(payload)

    if isinstance(payload, dict):
        for key in ("translated_data", "data", "items", "rows"):
            value = payload.get(key)
            if isinstance(value, list):
                return value
        return [payload]

    if isinstance(payload, list):
        return payload

    return []


def normalize_text(value: Any) -> str:
    return str(value).strip() if value not in (None, "") else ""


WORKFLOW_LOCATION_ORDER = [
    "右岸道路",
    "炸药库",
    "右岸砂石骨料加工系统",
    "右岸临时120拌合站",
    "右岸拌合系统",
    "右岸1#弃渣场",
    "右岸2#弃渣场",
    "右岸制砖厂兼钢筋加工厂",
    "右岸坝肩",
    "右岸基坑",
    "左岸2#弃渣场",
    "左岸EPC营地",
    "右岸施工营地",
    "22kV线路",
]

WORKFLOW_LOCATION_ALIASES = {
    "砂石系统": "右岸砂石骨料加工系统",
    "骨料系统": "右岸砂石骨料加工系统",
    "120拌合站": "右岸临时120拌合站",
    "拌合站": "右岸临时120拌合站",
    "右岸引航道": "右岸基坑",
    "引航道": "右岸基坑",
    "坝肩": "右岸坝肩",
    "22kV": "22kV线路",
}


# Override with corrected spoil-area mapping rules.
WORKFLOW_LOCATION_ORDER = [
    "右岸道路",
    "炸药库",
    "右岸砂石骨料加工系统",
    "右岸临时120拌合站",
    "右岸拌合系统",
    "右岸1#弃渣场",
    "右岸2#弃渣场",
    "右岸制砖厂兼钢筋加工厂",
    "右岸坝肩",
    "右岸基坑",
    "左岸1#弃渣场",
    "左岸2#弃渣场",
    "左岸EPC营地",
    "右岸施工营地",
    "22kV线路",
]

WORKFLOW_LOCATION_ALIASES = {
    "砂石系统": "右岸砂石骨料加工系统",
    "骨料系统": "右岸砂石骨料加工系统",
    "120拌合站": "右岸临时120拌合站",
    "拌合站": "右岸临时120拌合站",
    "右岸引航道": "右岸基坑",
    "引航道": "右岸基坑",
    "坝肩": "右岸坝肩",
    "22kV": "22kV线路",
}


def resolve_spoil_area_location(loc: str) -> str:
    if "弃渣场" not in loc:
        return ""
    side = "右岸"
    if "左岸" in loc:
        side = "左岸"
    elif "右岸" in loc:
        side = "右岸"
    number = "1"
    if "2#" in loc or "2号" in loc or "二号" in loc:
        number = "2"
    elif "1#" in loc or "1号" in loc or "一号" in loc:
        number = "1"
    return f"{side}{number}#弃渣场"


def normalize_location_with_workflow_logic(location: str) -> str:
    loc = location.strip()
    if not loc:
        return loc
    spoil_area = resolve_spoil_area_location(loc)
    if spoil_area:
        return spoil_area
    for standard in WORKFLOW_LOCATION_ORDER:
        if standard in loc or loc in standard:
            return standard
    for alias, target in WORKFLOW_LOCATION_ALIASES.items():
        if alias in loc:
            return target
    return loc


def should_keep_row(content: str) -> bool:
    text = content.strip()
    if not text:
        return False
    deny_keywords = ("盲沟", "未开工", "未施工", "无施工记录")
    return not any(word in text for word in deny_keywords)


def apply_workflow_matching_logic(rows: list[dict], *, locale: str) -> list[dict]:
    prepared: list[dict] = []
    for row in rows:
        content_key = "content" if locale == "cn" else "content_en"
        content = normalize_text(row.get(content_key))
        if not should_keep_row(content):
            continue
        new_row = dict(row)
        if locale == "cn":
            new_row["location"] = normalize_location_with_workflow_logic(normalize_text(row.get("location")))
        prepared.append(new_row)

    if locale == "cn":
        order_map = {name: idx for idx, name in enumerate(WORKFLOW_LOCATION_ORDER)}
        prepared.sort(key=lambda item: order_map.get(normalize_text(item.get("location")), 999))
    return prepared


def normalize_feishu_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (int, float, str)):
        return str(value)
    if isinstance(value, list):
        parts = [normalize_feishu_value(item) for item in value]
        return " ".join(part for part in parts if part)
    if isinstance(value, dict):
        for key in ("text", "name", "value"):
            if key in value:
                return normalize_feishu_value(value[key])
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def timestamp_to_date_strings(ms_or_s: float | int) -> list[str]:
    ts = float(ms_or_s)
    if ts > 1e12:
        ts /= 1000.0
    dt = datetime.fromtimestamp(ts)
    zh = f"{dt.year}年{dt.month}月{dt.day}日"
    return [dt.strftime("%Y-%m-%d"), zh, dt.strftime("%Y/%m/%d")]


def date_value_matches(field_value: Any, target_date_text: str) -> bool:
    if field_value is None or not target_date_text:
        return False
    tv = target_date_text.strip()
    if isinstance(field_value, (int, float)):
        return any(tv in s or s in tv for s in timestamp_to_date_strings(field_value))
    fv = normalize_feishu_value(field_value)
    return tv in fv or fv in tv


def pick_water_level_field_name(fields: dict[str, Any], preferred: str) -> str:
    candidates = [name.strip() for name in str(preferred).split(",") if name.strip()]
    candidates.extend(["水位高程", "水位", "water_level"])
    seen = set()
    for name in candidates:
        if name in seen:
            continue
        seen.add(name)
        if name in fields:
            return name
    return ""


def fetch_feishu_water_level(
    *,
    token: str,
    app_token: str,
    table_id: str,
    water_level_field: str = "水位高程",
    view_id: Optional[str] = None,
    date_field: Optional[str] = None,
    date_value: Optional[str] = None,
    app_id: Optional[str] = None,
    app_secret: Optional[str] = None,
) -> str:
    if not (app_token and table_id):
        return ""

    def get_token() -> str:
        if token:
            return token
        if not (app_id and app_secret):
            return ""
        resp = requests.post(
            "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal",
            json={"app_id": app_id, "app_secret": app_secret},
            timeout=10,
        )
        resp.raise_for_status()
        body = resp.json()
        if body.get("code") != 0:
            raise RuntimeError(f"Feishu auth failed: {body.get('msg', 'unknown error')}")
        return body.get("tenant_access_token", "")

    current_token = get_token()
    if not current_token:
        return ""

    url = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token}/tables/{table_id}/records"
    params: dict[str, Any] = {"page_size": 200}
    if view_id:
        params["view_id"] = view_id

    def read_items(bearer: str) -> list[dict]:
        resp = requests.get(
            url,
            headers={"Authorization": f"Bearer {bearer}"},
            params=params,
            timeout=10,
        )
        resp.raise_for_status()
        body = resp.json()
        if body.get("code") != 0:
            raise RuntimeError(f"Feishu records failed: {body.get('msg', 'unknown error')}")
        return body.get("data", {}).get("items", [])

    try:
        items = read_items(current_token)
    except Exception:
        if app_id and app_secret:
            current_token = get_token()
            items = read_items(current_token)
        else:
            raise

    if not items:
        return ""

    selected_fields = None
    if date_field and date_value:
        for item in items:
            fields = item.get("fields", {})
            if date_value_matches(fields.get(date_field), date_value):
                selected_fields = fields
                break

    if selected_fields is None:
        sort_field = date_field or ("观测日期" if "观测日期" in items[0].get("fields", {}) else "")
        if sort_field:
            items.sort(
                key=lambda x: x.get("fields", {}).get(sort_field) or 0,
                reverse=True,
            )
        selected_fields = items[0].get("fields", {})

    picked_field = pick_water_level_field_name(selected_fields, water_level_field)
    if not picked_field:
        return ""
    return normalize_feishu_value(selected_fields.get(picked_field))


def split_text_prefix(line_text: str) -> tuple[str, str]:
    for separator in ("：", ":"):
        if separator in line_text:
            split_index = line_text.index(separator) + 1
            return line_text[:split_index], line_text[split_index:]
    return line_text, ""


def normalize_seq_text(value: Any) -> str:
    text = normalize_text(value)
    if not text:
        return ""
    match = re.search(r"\d+", text)
    return match.group(0) if match else text


def build_report_rows(payload: Any, *, locale: str) -> list[dict[str, str]]:
    rows = parse_payload(payload)
    rows = apply_workflow_matching_logic(rows, locale=locale)
    prepared: list[dict[str, str]] = []
    current_seq = ""

    for item in rows:
        if not isinstance(item, dict):
            continue

        seq = normalize_seq_text(item.get("seq"))
        if seq:
            current_seq = seq
        if not current_seq:
            continue

        if locale == "cn":
            prepared.append(
                {
                    "seq": current_seq,
                    "location": normalize_text(item.get("location")),
                    "content": normalize_text(item.get("content")),
                    "quantity": normalize_text(item.get("quantity") or item.get("daily_qty")),
                    "remark": normalize_text(item.get("remark") or item.get("shift")),
                }
            )
        else:
            prepared.append(
                {
                    "seq": current_seq,
                    "location": normalize_text(item.get("location_en") or item.get("location")),
                    "content": normalize_text(item.get("content_en") or item.get("content")),
                    "quantity": normalize_text(
                        item.get("quantity_en") or item.get("daily_qty_en") or item.get("quantity") or item.get("daily_qty")
                    ),
                    "remark": normalize_text(item.get("remarks_en") or item.get("remark_en") or item.get("shift") or item.get("remark")),
                }
            )

    return prepared


def set_cell_text_with_format(cell, text: str, *, bold: bool = False):
    cell.text = ""
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    run = p.add_run(text)
    format_run_font(run, bold=bold)


def detect_template_groups(table, *, seq_col: int = 0) -> list[tuple[str, int, int]]:
    if not table.rows:
        return []

    seq_rows: list[int] = []
    for idx, row in enumerate(table.rows):
        if len(row.cells) <= seq_col:
            continue
        if normalize_seq_text(row.cells[seq_col].text):
            seq_rows.append(idx)

    if not seq_rows:
        return []

    start_row = seq_rows[0]
    prev_seq = normalize_seq_text(table.rows[start_row].cells[seq_col].text)
    group_start = start_row
    groups: list[tuple[str, int, int]] = []

    for i in range(start_row + 1, len(table.rows)):
        if len(table.rows[i].cells) <= seq_col:
            continue
        current_seq = normalize_seq_text(table.rows[i].cells[seq_col].text)
        if current_seq and current_seq != prev_seq:
            groups.append((prev_seq, group_start, i - 1))
            group_start = i
            prev_seq = current_seq

    groups.append((prev_seq, group_start, len(table.rows) - 1))
    return [(seq, start, end) for seq, start, end in groups if seq]


def render_rows_to_grouped_table(table, payload: Any, *, locale: str) -> bool:
    if not table.rows:
        return False

    cols_count = len(table.rows[0].cells)
    if cols_count < 4:
        return False

    seq_col = 0
    location_col = 1
    if cols_count >= 6:
        content_col, qty_col, remark_col = 3, 4, 5
    else:
        content_col, qty_col, remark_col = 2, 3, 4

    groups = detect_template_groups(table, seq_col=seq_col)
    if not groups:
        return False

    rows = build_report_rows(payload, locale=locale)
    if not rows:
        return False

    grouped_input: dict[str, list[dict[str, str]]] = {}
    for row in rows:
        seq = normalize_seq_text(row.get("seq"))
        if not seq:
            continue
        grouped_input.setdefault(seq, []).append(row)

    matched_group = False
    for seq, start, end in groups:
        items = grouped_input.get(seq, [])
        if not items:
            continue

        matched_group = True
        first = items[0]
        if start < len(table.rows) and len(table.rows[start].cells) > seq_col:
            set_cell_text_with_format(table.rows[start].cells[seq_col], seq)
        if start < len(table.rows) and len(table.rows[start].cells) > location_col:
            set_cell_text_with_format(table.rows[start].cells[location_col], first.get("location", ""))

        span = end - start + 1
        for offset in range(span):
            row_idx = start + offset
            if row_idx >= len(table.rows):
                break
            target_cells = table.rows[row_idx].cells
            if offset < len(items):
                data = items[offset]
                if len(target_cells) > content_col:
                    set_cell_text_with_format(target_cells[content_col], data.get("content", ""))
                if len(target_cells) > qty_col:
                    set_cell_text_with_format(target_cells[qty_col], data.get("quantity", ""))
                if remark_col >= 0 and len(target_cells) > remark_col:
                    set_cell_text_with_format(target_cells[remark_col], data.get("remark", ""))
            else:
                if len(target_cells) > content_col:
                    set_cell_text_with_format(target_cells[content_col], "")
                if len(target_cells) > qty_col:
                    set_cell_text_with_format(target_cells[qty_col], "")
                if remark_col >= 0 and len(target_cells) > remark_col:
                    set_cell_text_with_format(target_cells[remark_col], "")

    return matched_group


def build_report_blocks(payload: Any, *, locale: str) -> list[dict[str, Any]]:
    """把原始数据整理成更接近 Word 报表布局的分组块。"""
    rows = parse_payload(payload)
    rows = apply_workflow_matching_logic(rows, locale=locale)
    blocks: list[dict[str, Any]] = []
    current_key: Any = object()

    for item in rows:
        if not isinstance(item, dict):
            continue

        seq = item.get("seq")
        if locale == "cn":
            heading = normalize_text(item.get("location"))
            title = f"{seq}、{heading}" if seq not in (None, "") and heading else heading
            detail = " ".join(part for part in (normalize_text(item.get("content")), normalize_text(item.get("quantity"))) if part)
            group_key = (seq, heading)
        else:
            heading = normalize_text(item.get("location_en") or item.get("location"))
            title = f"{seq}. {heading}" if seq not in (None, "") and heading else heading
            detail_parts = [
                normalize_text(item.get("content_en")),
                normalize_text(item.get("quantity_en")),
                normalize_text(item.get("remarks_en") or item.get("shift")),
            ]
            detail = " ".join(part for part in detail_parts if part)
            group_key = (seq, heading)

        if not blocks or group_key != current_key:
            blocks.append({"title": title, "lines": []})
            current_key = group_key

        if detail:
            blocks[-1]["lines"].append(detail)

    return blocks


def render_text_line(
    paragraph,
    line_text: str,
    *,
    indent: int = 24,
    bold_prefix: bool = False,
    bold: bool = False,
):
    paragraph.paragraph_format.first_line_indent = Pt(indent)
    if bold_prefix:
        prefix, suffix = split_text_prefix(line_text)
        run_prefix = paragraph.add_run(prefix)
        format_run_font(run_prefix, bold=True)
        if suffix:
            run_suffix = paragraph.add_run(suffix)
            format_run_font(run_suffix, bold=False)
        return

    run = paragraph.add_run(line_text)
    format_run_font(run, bold=bold)


def append_report_block(cell, title: str, lines: list[str]):
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    title_run = p.add_run(title)
    format_run_font(title_run, bold=True)

    for line in lines:
        paragraph = cell.add_paragraph()
        render_text_line(paragraph, line, indent=24)


def render_report_blocks_to_cell(cell, payload: Any, *, locale: str):
    cell.text = ""
    for block in build_report_blocks(payload, locale=locale):
        title = normalize_text(block.get("title"))
        if not title:
            continue
        lines = [normalize_text(line) for line in block.get("lines", []) if normalize_text(line)]
        append_report_block(cell, title, lines)


def render_report_blocks_to_document(
    doc: Document,
    payload: Any,
    *,
    locale: str,
    table_index: int = 0,
    row_index: int = 4,
    col_index: int = 2,
):
    if doc.tables and len(doc.tables) > table_index:
        table = doc.tables[table_index]
        if render_rows_to_grouped_table(table, payload, locale=locale):
            return
        if len(table.rows) > row_index and len(table.rows[row_index].cells) > col_index:
            render_report_blocks_to_cell(table.cell(row_index, col_index), payload, locale=locale)
            return

    for block in build_report_blocks(payload, locale=locale):
        title = normalize_text(block.get("title"))
        if title:
            paragraph = doc.add_paragraph()
            render_text_line(paragraph, title, indent=0, bold=True)
        for line in block.get("lines", []):
            paragraph = doc.add_paragraph()
            render_text_line(paragraph, normalize_text(line), indent=24)


def write_lines_to_cell(cell, content: str):
    cell.text = ""
    for line in content.splitlines():
        process_and_add_line(cell, line)


def populate_template(
    doc: Document,
    content: str,
    table_index: int = 0,
    row_index: int = 4,
    col_index: int = 2,
):
    """尽量把内容写进目标表格单元格；如果没有目标表格，则退化成普通段落。"""
    if doc.tables and len(doc.tables) > table_index:
        table = doc.tables[table_index]
        if len(table.rows) > row_index and len(table.rows[row_index].cells) > col_index:
            write_lines_to_cell(table.cell(row_index, col_index), content)
            return

    # 兼容没有模板或模板结构变化的情况，避免接口直接失败。
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        run = p.add_run(line)
        format_run_font(run, bold=False)


def update_header_date_weather(
    doc: Document,
    *,
    locale: str = "cn",
    feishu_token: Optional[str] = None,
    feishu_app_token: Optional[str] = None,
    feishu_table_id: Optional[str] = None,
    feishu_view_id: Optional[str] = None,
    feishu_water_level_field: str = "\u6c34\u4f4d\u9ad8\u7a0b",
    feishu_date_field: Optional[str] = None,
    feishu_date_value: Optional[str] = None,
    feishu_app_id: Optional[str] = None,
    feishu_app_secret: Optional[str] = None,
    warnings: Optional[List[Dict[str, str]]] = None,
    warning_target: str = "document",
) -> dict[str, str]:
    now = datetime.now()
    date_str = f"{now.year}年{now.month}月{now.day}日" if locale == "cn" else now.strftime("%Y-%m-%d")
    weather = "晴" if locale == "cn" else "Sunny"
    temp = "20℃~30℃" if locale == "cn" else "20C~30C"
    water_level = ""
    water_level_status = "ok"
    water_level_error = ""

    try:
        water_level = fetch_feishu_water_level(
            token=feishu_token or "",
            app_token=feishu_app_token or "",
            table_id=feishu_table_id or "",
            water_level_field=feishu_water_level_field,
            view_id=feishu_view_id,
            date_field=feishu_date_field,
            date_value=feishu_date_value,
            app_id=feishu_app_id,
            app_secret=feishu_app_secret,
        )
        if not water_level:
            water_level_error = "Feishu returned empty water level."
    except Exception as exc:
        water_level_error = str(exc)
        water_level = ""

    if water_level_error:
        water_level_status = "failed"
        if warnings is not None:
            warnings.append(
                {
                    "code": "FEISHU_WATER_LEVEL_FAILED",
                    "target": warning_target,
                    "message": water_level_error,
                }
            )

    if doc.tables and len(doc.tables[0].rows) > 0:
        cells = doc.tables[0].rows[0].cells
        if len(cells) > 0:
            cells[0].text = ""
            run = cells[0].paragraphs[0].add_run(date_str)
            format_run_font(run)
        if len(cells) > 1:
            weather_str = (
                f"天气：{weather}                气温：{temp}"
                if locale == "cn"
                else f"Weather: {weather}                Temperature: {temp}"
            )
            cells[1].text = ""
            run = cells[1].paragraphs[0].add_run(weather_str)
            format_run_font(run)
        if len(cells) > 3:
            cells[3].text = ""
            run = cells[3].paragraphs[0].add_run(water_level or "")
            format_run_font(run)

    return {
        "date": date_str,
        "weather": weather,
        "temp": temp,
        "water_level": water_level,
        "water_level_status": water_level_status,
    }

def process_and_add_line(cell, line_text):
    """
    智能处理每一行的格式：缩进、加粗、分割
    """
    line_text = line_text.strip()
    if not line_text: return

    # 创建新段落（注意：不使用 add_run("\n") 而是 add_paragraph 以便单独控制每一行的缩进）
    # 如果是单元格刚清空后的第一个段落，直接使用，否则新建
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    # --- 1. 规则匹配 ---
    
    # 规则A：大标题 (例如 "1、右岸施工营地")
    # 特征：数字开头 + 顿号或点
    if re.match(r"^\d+[、\.]", line_text):
        p.paragraph_format.first_line_indent = Pt(0) # 【关键】强制不缩进
        run = p.add_run(line_text)
        format_run_font(run, bold=True) # 大标题整行加粗
        return

    # 规则B：统计项 (例如 "人员投入：...")
    # 特征：包含特定关键词
    keywords = ["人员投入", "设备投入", "累计工程量"]
    hit_keyword = None
    for kw in keywords:
        if kw in line_text:
            hit_keyword = kw
            break
    
    if hit_keyword:
        p.paragraph_format.first_line_indent = Pt(24) # 【关键】强制缩进 2 字符
        
        # 【局部加粗逻辑】
        # 将文本切分为两部分：关键词前缀(加粗) + 剩余内容(不加粗)
        # 例如 "人员投入：张三" -> "人员投入：" (粗) + " 张三" (细)
        
        # 尝试找到冒号的位置
        split_index = -1
        if "：" in line_text:
            split_index = line_text.index("：") + 1
        elif ":" in line_text:
            split_index = line_text.index(":") + 1
        else:
            # 如果没有冒号，就只加粗关键词本身
            split_index = line_text.index(hit_keyword) + len(hit_keyword)
            
        prefix = line_text[:split_index]
        content = line_text[split_index:]
        
        # 写入前缀（加粗）
        run1 = p.add_run(prefix)
        format_run_font(run1, bold=True)
        
        # 写入内容（正常）
        run2 = p.add_run(content)
        format_run_font(run2, bold=False)
        return

    # 规则C：子标题 / 具体内容 (例如 "(1) 场地精平")
    # 特征：以 (数字) 或 （数字） 开头
    if re.match(r"^[\(（]\d+[\)）]", line_text):
        p.paragraph_format.first_line_indent = Pt(24) # 【关键】强制缩进 2 字符
        run = p.add_run(line_text)
        format_run_font(run, bold=False) # 内容不加粗
        return

    # 规则D：其他普通文本
    # 默认缩进2字符（因为通常是正文延续），或者0？
    # 根据你的截图，如果不符合上述规则，通常是正文描述，建议缩进2字符对齐
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(line_text)
    format_run_font(run, bold=False)

# ---------------- 辅助函数 ----------------

def find_target_table(doc: Document, index: int) -> Optional[Any]:
    if 0 <= index < len(doc.tables):
        return doc.tables[index]
    return None

def update_table_row(table, row_name: str, today: str, total: str):
    """表格行更新逻辑"""
    if not table.rows: return
    
    name_col = 1
    cols_count = len(table.rows[0].cells)
    today_col = 4 if cols_count > 4 else cols_count - 2
    total_col = 5 if cols_count > 5 else cols_count - 1
    
    for row in table.rows:
        if len(row.cells) <= max(name_col, today_col, total_col): continue
        cell_text = row.cells[name_col].text.strip()
        if row_name in cell_text: 
            # 填入数字时也应用字体规范
            if today and today != "-":
                cell = row.cells[today_col]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(today))
                format_run_font(run)
            if total and total != "-":
                cell = row.cells[total_col]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(total))
                format_run_font(run)
            return

# ---------------- 模型定义 ----------------

class FillTemplateRequest(BaseModel):
    template_base64: str
    content: str
    table_index: int = 0
    row_index: int = 4
    col_index: int = 2
    update_date_weather: bool = False
    upload_to_feishu: bool = False
    feishu_token: Optional[str] = None
    feishu_app_token: Optional[str] = None
    feishu_table_id: Optional[str] = None
    feishu_view_id: Optional[str] = None
    feishu_water_level_field: str = "\u6c34\u4f4d\u9ad8\u7a0b"
    feishu_date_field: Optional[str] = None
    feishu_date_value: Optional[str] = None
    feishu_app_id: Optional[str] = None
    feishu_app_secret: Optional[str] = None

class UpdateDateWeatherRequest(BaseModel):
    document_base64: str
    feishu_token: Optional[str] = None
    feishu_app_token: Optional[str] = None
    feishu_table_id: Optional[str] = None
    feishu_view_id: Optional[str] = None
    feishu_water_level_field: str = "\u6c34\u4f4d\u9ad8\u7a0b"
    feishu_date_field: Optional[str] = None
    feishu_date_value: Optional[str] = None
    feishu_app_id: Optional[str] = None
    feishu_app_secret: Optional[str] = None

class UpdatePersonnelRequest(BaseModel):
    document_base64: str
    personnel_text: str 
    feishu_token: Optional[str] = None

class AppendixTableData(BaseModel):
    table_index: int
    row_name: str
    today_qty: str
    total_qty: str

class UpdateAppendixRequest(BaseModel):
    document_base64: str
    data: List[AppendixTableData]
    feishu_token: Optional[str] = None


class GenerateFromTemplateRequest(BaseModel):
    chinese_data: Any
    english_data: Any
    cn_template_base64: Optional[str] = None
    en_template_base64: Optional[str] = None
    update_date_weather: bool = False
    cn_table_index: int = 0
    cn_row_index: int = 4
    cn_col_index: int = 2
    en_table_index: int = 0
    en_row_index: int = 4
    en_col_index: int = 2
    feishu_token: Optional[str] = None
    feishu_app_token: Optional[str] = None
    feishu_table_id: Optional[str] = None
    feishu_view_id: Optional[str] = None
    feishu_water_level_field: str = "\u6c34\u4f4d\u9ad8\u7a0b"
    feishu_date_field: Optional[str] = None
    feishu_date_value: Optional[str] = None
    feishu_app_id: Optional[str] = None
    feishu_app_secret: Optional[str] = None

# ---------------- API 接口实现 ----------------

@app.get("/health")
async def health():
    return {"status": "ok"}


@app.post("/fill-template")
async def fill_template(req: FillTemplateRequest):
    try:
        def mutator(doc: Document):
            populate_template(doc, req.content, req.table_index, req.row_index, req.col_index)
            extra = {}
            if req.update_date_weather:
                extra["weather_info"] = update_header_date_weather(
                    doc,
                    feishu_token=req.feishu_token,
                    feishu_app_token=req.feishu_app_token,
                    feishu_table_id=req.feishu_table_id,
                    feishu_view_id=req.feishu_view_id,
                    feishu_water_level_field=req.feishu_water_level_field,
                    feishu_date_field=req.feishu_date_field,
                    feishu_date_value=req.feishu_date_value,
                    feishu_app_id=req.feishu_app_id,
                    feishu_app_secret=req.feishu_app_secret,
                )
            return extra

        return update_document(req.template_base64, mutator)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "message": str(e)}

@app.post("/update-date-weather")
async def update_date_weather(req: UpdateDateWeatherRequest):
    try:
        return update_document(
            req.document_base64,
            lambda doc: {
                "weather_info": update_header_date_weather(
                    doc,
                    feishu_token=req.feishu_token,
                    feishu_app_token=req.feishu_app_token,
                    feishu_table_id=req.feishu_table_id,
                    feishu_view_id=req.feishu_view_id,
                    feishu_water_level_field=req.feishu_water_level_field,
                    feishu_date_field=req.feishu_date_field,
                    feishu_date_value=req.feishu_date_value,
                    feishu_app_id=req.feishu_app_id,
                    feishu_app_secret=req.feishu_app_secret,
                )
            },
        )
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.post("/update-personnel-stats")
async def update_personnel_stats(req: UpdatePersonnelRequest):
    try:
        def mutator(doc: Document):
            p = doc.add_paragraph()
            run = p.add_run(req.personnel_text.strip())
            format_run_font(run)

        return update_document(req.document_base64, mutator)
    except Exception as e:
        return {"success": False, "message": str(e)}

@app.post("/update-appendix-tables")
async def update_appendix_tables(req: UpdateAppendixRequest):
    try:
        def mutator(doc: Document):
            for item in req.data:
                target_table = find_target_table(doc, item.table_index)
                if target_table:
                    update_table_row(target_table, item.row_name, item.today_qty, item.total_qty)

        return update_document(req.document_base64, mutator)
    except Exception as e:
        return {"success": False, "message": str(e)}


@app.post("/generate-from-template")
async def generate_from_template(req: GenerateFromTemplateRequest):
    try:
        cn_doc = decode_document(req.cn_template_base64) if req.cn_template_base64 else Document()
        en_doc = decode_document(req.en_template_base64) if req.en_template_base64 else Document()
        warnings: list[dict[str, str]] = []

        render_report_blocks_to_document(
            cn_doc,
            req.chinese_data,
            locale="cn",
            table_index=req.cn_table_index,
            row_index=req.cn_row_index,
            col_index=req.cn_col_index,
        )
        render_report_blocks_to_document(
            en_doc,
            req.english_data,
            locale="en",
            table_index=req.en_table_index,
            row_index=req.en_row_index,
            col_index=req.en_col_index,
        )

        result = {
            "success": True,
            "cn_document_base64": encode_document(cn_doc),
            "en_document_base64": encode_document(en_doc),
        }

        if req.update_date_weather:
            result["weather_info"] = update_header_date_weather(
                cn_doc,
                locale="cn",
                feishu_token=req.feishu_token,
                feishu_app_token=req.feishu_app_token,
                feishu_table_id=req.feishu_table_id,
                feishu_view_id=req.feishu_view_id,
                feishu_water_level_field=req.feishu_water_level_field,
                feishu_date_field=req.feishu_date_field,
                feishu_date_value=req.feishu_date_value,
                feishu_app_id=req.feishu_app_id,
                feishu_app_secret=req.feishu_app_secret,
                warnings=warnings,
                warning_target="cn_document",
            )
            update_header_date_weather(
                en_doc,
                locale="en",
                feishu_token=req.feishu_token,
                feishu_app_token=req.feishu_app_token,
                feishu_table_id=req.feishu_table_id,
                feishu_view_id=req.feishu_view_id,
                feishu_water_level_field=req.feishu_water_level_field,
                feishu_date_field=req.feishu_date_field,
                feishu_date_value=req.feishu_date_value,
                feishu_app_id=req.feishu_app_id,
                feishu_app_secret=req.feishu_app_secret,
                warnings=warnings,
                warning_target="en_document",
            )
            result["cn_document_base64"] = encode_document(cn_doc)
            result["en_document_base64"] = encode_document(en_doc)
            if warnings:
                result["warnings"] = warnings

        return result
    except Exception as e:
        return {"success": False, "message": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
